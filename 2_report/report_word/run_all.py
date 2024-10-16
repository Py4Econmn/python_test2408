# pip install python-docx
from docx import Document

dir = 'report_word/'

doc = Document(dir + 'report.docx')

# for para in doc.paragraphs:
#     print("-----------------------")
#     print(para.text)

# print("\nTables in the document:")
# for table in doc.tables:
#     for row in table.rows:
#         for cell in row.cells:
#             print(cell.text, end=' | ')
#         print()  # new line after each row


from docx import Document

def replace_text_in_paragraph(paragraph, replacements):
    for placeholder, value in replacements.items():
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, value)

def replace_placeholders(doc, replacements):
    # Replace placeholders in paragraphs
    for para in doc.paragraphs:
        replace_text_in_paragraph(para, replacements)
    
    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_text_in_paragraph(para, replacements)

# Example usage
replacements = {
    '{{DATE}}': '2024 оны 2 дугаар улирлын',
    '{{free_lag2}}': '155'
}


replace_placeholders(doc, replacements)
doc.save(dir + 'filled_report.docx')



# Add a new paragraph
# doc.add_paragraph('This is a new paragraph added by Python.')

# Save the document
# doc.save('automated_report.docx')

paras = doc.paragraphs
for para in doc.paragraphs:
    for run in para.runs:
        print('-----------------------')
        print(run.text)



for run in paras[4].runs:
    print('-----------------------')
    print(run.text)

