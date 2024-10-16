from docx import Document

def replace_text_in_run(run, placeholder, value):
    if placeholder in run.text:
        run.text = run.text.replace(placeholder, value)

def replace_placeholders(doc, replacements):
    # Replace placeholders in paragraphs while preserving formatting
    for para in doc.paragraphs:
        for run in para.runs:
            for placeholder, value in replacements.items():
                replace_text_in_run(run, placeholder, value)
    
    # Replace placeholders in tables while preserving formatting
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for placeholder, value in replacements.items():
                            replace_text_in_run(run, placeholder, value)

# Example usage
replacements = {
    '{{pl_date}}': '2024-08-29',
    '{{SUMMARY}}': 'This is the summary of the report.',
    '{{HEADER1}}': '125',
    '{{HEADER2}}': 'Header 2',
    '{{DATA1}}': 'Data 1',
    '{{DATA2}}': 'Data 2'
}

doc = Document('report.docx')
replace_placeholders(doc, replacements)
doc.save('filled_report.docx')
