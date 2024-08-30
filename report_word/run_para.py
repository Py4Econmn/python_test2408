from docx import Document
from docx.shared import Pt

def set_font_for_paragraph(paragraph, font_name='Roboto Light', font_size=11):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)

def set_font_for_table_cell(cell, font_name='Roboto Light', font_size=9):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)

def replace_text_in_paragraph(paragraph, replacements):
    original_text = paragraph.text
    for placeholder, value in replacements.items():
        if placeholder in original_text:
            # Replace the text
            paragraph.text = original_text.replace(placeholder, value)
            # Set font to Roboto Light if replacement occurred
            set_font_for_paragraph(paragraph)
            break  # Only need to set the font once per paragraph

def replace_text_in_table(cell, replacements):
    for para in cell.paragraphs:
        original_text = para.text
        for placeholder, value in replacements.items():
            if placeholder in original_text:
                # Replace the text
                para.text = original_text.replace(placeholder, value)
                # Set font to Roboto Light if replacement occurred
                set_font_for_table_cell(cell)
                break  # Only need to set the font once per paragraph

def replace_placeholders(doc, replacements):
    # Replace placeholders in paragraphs
    for para in doc.paragraphs:
        replace_text_in_paragraph(para, replacements)
    
    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_table(cell, replacements)

# Example usage
replacements = {
    '{{DATE}}': '2024-08-29',
    '{{SUMMARY}}': 'This is the summary of the report.',
    '{{HEADER1}}': '53.6',
    '{{HEADER2}}': 'Header 2',
    '{{DATA1}}': 'Data 1',
    '{{DATA2}}': 'Data 2'
}

doc = Document('report.docx')
replace_placeholders(doc, replacements)
doc.save('filled_report.docx')
