from docx import Document
from docx.shared import Inches

def replace_image_by_order(doc, figure_order, new_image_path, width_in_inches=3.0, height_in_inches=3):
    """
    Replace an image by its order in the document with a new image.

    :param doc: The Document object to modify.
    :param figure_order: The order of the figure to replace (1-based index).
    :param new_image_path: Path to the new image file.
    :param width_in_inches: Width of the new image in inches.
    :param height_in_inches: Height of the new image in inches (optional).
    """
    figure_count = 0

    # Process paragraphs
    for para in doc.paragraphs:
        for shape in para._element.xpath(".//w:drawing"):
            figure_count += 1
            if figure_count == figure_order:
                # Remove the existing image
                para._element.remove(shape)
                # Insert the new image
                para.add_run().add_picture(new_image_path, width=Inches(width_in_inches), height=Inches(height_in_inches))
                return

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for shape in para._element.xpath(".//w:drawing"):
                        figure_count += 1
                        if figure_count == figure_order:
                            # Remove the existing image
                            para._element.remove(shape)
                            # Insert the new image
                            para.add_run().add_picture(new_image_path, width=Inches(width_in_inches), height=Inches(height_in_inches))
                            return

# Example usage
doc = Document('report.docx')

# Path to the new image file
new_image_path = 'new_figure.png'

# Replace the second image in the document
replace_image_by_order(doc, 2, new_image_path, width_in_inches=4.0)

# Save the document
doc.save('report_with_replaced_image.docx')
