from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn

def replace_placeholder_with_image(doc, placeholder, image_path, width_in_inches=3.0, height_in_inches=2):
    """
    Replace a placeholder with an image and set text wrapping properties.

    :param doc: The Document object to modify.
    :param placeholder: The placeholder text to be replaced.
    :param image_path: Path to the image file.
    :param width_in_inches: Width of the image in inches.
    :param height_in_inches: Height of the image in inches (optional).
    """
    for para in doc.paragraphs:
        if placeholder in para.text:
            # Remove placeholder text
            para.text = para.text.replace(placeholder, '')
            # Add the picture
            run = para.add_run()
            picture = run.add_picture(image_path, width=Inches(width_in_inches), height=Inches(height_in_inches))

            # Get the XML element for the picture
            inline_shape = run._element.xpath('.//a:blip')[0].getparent()
            # Set text wrapping properties
            inline_shape.getparent().set(qn('w:wrap'), 'square')  # or 'tight', 'behind', etc.
            return

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if placeholder in para.text:
                        # Remove placeholder text
                        para.text = para.text.replace(placeholder, '')
                        # Add the picture
                        run = para.add_run()
                        picture = run.add_picture(image_path, width=Inches(width_in_inches), height=Inches(height_in_inches))

                        # Get the XML element for the picture
                        inline_shape = run._element.xpath('.//a:blip')[0].getparent()
                        # Set text wrapping properties
                        inline_shape.getparent().set(qn('w:wrap'), 'behind')  # or 'tight', 'behind', etc.
                        inline_shape.getparent().set(qn('w:pos'), 'left')
                        return

# Example usage
doc = Document('report.docx')

# Path to the image file
image_path = 'figure_bop.png'

# Replace placeholder with image and set text wrapping
replace_placeholder_with_image(doc, '{{PLACEHOLDER}}', image_path, width_in_inches=2.0)

# Save the document
doc.save('report_with_wrapped_image.docx')
