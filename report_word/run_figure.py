from docx import Document
from docx.shared import Inches

def replace_placeholder_with_image(doc, placeholder, image_path, width_in_inches=4.0, height_in_inches=2):
    """
    Replace a placeholder with an image in the document.

    :param doc: The Document object to modify.
    :param placeholder: The placeholder text to be replaced.
    :param image_path: Path to the image file.
    :param width_in_inches: Width of the image in inches.
    :param height_in_inches: Height of the image in inches (optional).
    """
    for para in doc.paragraphs:
        if placeholder in para.text:
            # Remove placeholder text
            para.text = para.text.replace(placeholder, '')
            # Add picture in place of the placeholder
            para.add_run().add_picture(image_path, width=Inches(width_in_inches), height=Inches(height_in_inches))
            break  # Assuming there is only one placeholder of this type

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if placeholder in para.text:
                        # Remove placeholder text
                        para.text = para.text.replace(placeholder, '')
                        # Add picture in place of the placeholder
                        para.add_run().add_picture(image_path, width=Inches(width_in_inches), height=Inches(height_in_inches))
                        break  # Assuming there is only one placeholder of this type

dir = 'report_word/'
# Example usage
doc = Document(dir + 'report.docx')

# Path to the image file
image_path = dir + 'current_account_plot.png'

# Replace placeholder with image
replace_placeholder_with_image(doc, '{{Figure1}}', image_path, width_in_inches=7.0)

# Save the document
doc.save(dir + 'report_with_image.docx')
